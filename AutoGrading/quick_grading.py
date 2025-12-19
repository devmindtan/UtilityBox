import json
import re
from pathlib import Path

import filetype
import fitz
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document

# from dotenv import load_dotenv
from google import genai
from google.genai import types

# load_dotenv()


# pip install filetype pymupdf pandas beautifulsoup4 python-docx google-genai python-dotenv

FOLDER_PATH = (
    r"/home/devmindtan/Documents/Code/UtilityBox/AutoGrading/Assignments/Lab9/"
)
TOPIC_FILE = FOLDER_PATH + r"/Topic/Bai_tap_string_p2.docx"
API_KEY = "AIzaSyAFTB2zxZs23M9QNeyAJszpLy8_dcvPJE8"
MODEL_NAME = "gemma-3-12b"

# ========== INIT CLIENT ==========
client = genai.Client(api_key=API_KEY)


# Xác định kiểu của ảnh
def detect_mine_type(image_data):
    mime_type = ""

    kind = filetype.guess(image_data)
    if kind:
        mime_type = kind.mime
    else:
        mime_type = "image/png"

    return mime_type


# Xử lí đầu vào
def safe_json_loads(text):
    # Cắt bỏ mọi thứ ngoài JSON thật
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        clean = match.group(0).strip()
        return json.loads(clean)
    else:
        raise ValueError("Không tìm thấy JSON hợp lệ trong phản hồi")


def read_doc_content():
    folder_path = Path(FOLDER_PATH)
    data = []

    for doc_file in folder_path.glob("*.docx"):
        doc = Document(str(doc_file))
        rels = doc.part.rels

        # ======= Đọc text =======
        all_text = []
        for para in doc.paragraphs:
            all_text.append(para.text.strip())
        text_output = "\n".join(all_text)

        # ======= Đọc ảnh =======
        image_parts = []
        for rel in rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                image_parts.append(
                    types.Part.from_bytes(
                        data=image_data, mime_type=detect_mine_type(image_data)
                    )
                )

        data.append(
            {
                "file_name": doc_file.stem.strip(),
                "text": text_output,
                "images": image_parts,
            }
        )

    return data


def read_pdf_content():
    folder_path = Path(FOLDER_PATH)
    data = []

    for pdf_file in folder_path.glob("*.pdf"):
        pdf = fitz.open(pdf_file)
        all_text = []
        image_parts = []

        for page in pdf:
            # Lấy text
            all_text.append(page.get_text())

            # Lấy ảnh
            for img in page.get_images(full=True):
                xref = img[0]
                image_data = pdf.extract_image(xref)["image"]
                image_parts.append(
                    types.Part.from_bytes(
                        data=image_data, mime_type=detect_mine_type(image_data)
                    )
                )

        text_output = "\n".join(all_text)

        data.append(
            {
                "file_name": pdf_file.stem.strip(),
                "text": text_output,
                "images": image_parts,
            }
        )

    return data


def read_topic_content():
    path = Path(TOPIC_FILE)
    if not path.exists():
        raise FileNotFoundError(f"File đề bài không tồn tại: {TOPIC_FILE}")

    # ===== HTML =====
    if path.suffix.lower() == ".html":
        with open(path, "r", encoding="utf-8") as f:
            html = f.read()
        soup = BeautifulSoup(html, "html.parser")
        return {"content": soup.get_text().strip(), "image_content": []}

    # ===== DOCX =====
    elif path.suffix.lower() == ".docx":
        doc = Document(str(path))
        content = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
        rels = doc.part.rels.values()
        image_parts = []
        for rel in rels:
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                image_parts.append(
                    types.Part.from_bytes(
                        data=image_data, mime_type=detect_mine_type(image_data)
                    )
                )
        return {"content": content, "image_content": image_parts}

    else:
        raise ValueError("File đề bài phải là HTML hoặc DOCX")


def read_all_content():
    all_data = []

    # DOCX
    all_data.extend(read_doc_content())

    # PDF
    all_data.extend(read_pdf_content())

    return all_data


# ========== TẠO PROMPT ==========

RUBRIC_1 = """
    # Tiêu chí Chấm điểm Chi tiết (Thang 10)
    
    1. **Code (Tổng: 7 điểm)**
        a. **Hoàn thành:** Hoàn thành đầy đủ các phần code theo yêu cầu của đề bài. (Tối đa 4 điểm)
        b. **Chạy được:** Code phải chạy được và không có lỗi nghiêm trọng. (Tối đa 2 điểm)
        c. **Chất lượng:** Code rõ ràng, dễ đọc, có tuân thủ quy tắc lập trình cơ bản. (Tối đa 1 điểm)

    2. **Giải thích (Tổng: 3 điểm)**
        a. **Ý tưởng:** Trình bày được **ý tưởng chung** và **nguyên lý hoạt động** của đoạn code mẫu. (Tối đa 2 điểm)
        b. **Tuân thủ Yêu cầu:** Không mô tả từng dòng code. (Tối đa 1 điểm)
        
    - **Yêu cầu chung:** Điểm thưởng nếu có minh chứng kết quả chạy được (Ảnh chụp, output...).
    """
RUBRIC_2 = f"""
    # Tiêu chí Chấm điểm Chi tiết (Thang 10)
    
    1. **Nộp code (Tổng: 4 điểm)**
        a. **Đầy đủ:** Có đầy đủ file/code theo yêu cầu của đề bài (ví dụ: đủ 4 bài). (Tối đa 4 điểm)
        
    2. **Giải thích ý tưởng (Tổng: 5 điểm)**
        a. **Ý tưởng chung:** Trình bày đúng và rõ ràng **ý tưởng/nguyên lý** hoạt động của đoạn code/bài tập mẫu. (Tối đa 4 điểm)
        b. **Tuân thủ Yêu cầu:** KHÔNG mô tả từng dòng code; *Giải thích ý tưởng phải được trình bày trước khi code*. (Tối đa 1 điểm)
        
    3. **Kết quả chạy được (Tổng: 1 điểm)**
        a. **Minh chứng:** Có minh chứng (ảnh chụp, output,...) cho thấy code chạy được thành công và không lỗi. (Tối đa 1 điểm)
    """


def create_prompt(file_name, content):
    return f"""
Bạn là giảng viên đại học chấm bài sinh viên. Nhiệm vụ của bạn là đánh giá bài làm theo **Nguyên tắc Tính Nhất Quán Cao (High Consistency)**.

## Nguyên tắc:
**BẮT BUỘC** phải đánh giá và cộng điểm theo **TỪNG TIÊU CHÍ CON** trong **RUBRIC CHI TIẾT** dưới đây trước khi đưa ra Điểm số cuối cùng.

## Nhiệm vụ của bạn:
1. Thực hiện đánh giá chi tiết theo **RUBRIC CHI TIẾT**.
2. Tính **tổng điểm** dựa trên điểm của từng tiêu chí con (tối đa 10 điểm).
3. Viết nhận xét ngắn gọn (1-2 câu, dưới 20 chữ), xưng "em", giọng chuyên nghiệp.
4. Trả về **duy nhất một khối văn bản (text block)** theo định dạng sau:
    Tên file: {file_name}
    Điểm số: <điểm thang 10>
    Nhận xét: <Nhận xét ngắn gọn 1-2 câu>
---

## Thông tin chấm:
**Đề bài:**
{read_doc_content()}

**Rubric:**
{RUBRIC_2}

**Nội dung bài làm:**
{content}
"""


# ========== GỌI GEMINI ==========
def genemi_call(prompt_text, images):
    response = client.models.generate_content(
        model=MODEL_NAME, contents=[types.Part.from_text(text=prompt_text), *images]
    )

    return response.text


# ========== CHẤM ĐIỂM HÀNG LOẠT ==========


def grading():
    all_res = []
    for doc in read_all_content():

        try:
            print(f"📄 {doc['file_name']} — Đã đọc {len(doc['images'])} ảnh")

            file_name = doc["file_name"]
            content = doc["text"]
            images = doc["images"]

            # Tạo prompt
            prompt_text = create_prompt(file_name, content)

            # Gọi Gemini
            res = genemi_call(prompt_text, images)

            # Parse JSON
            # all_res.append(safe_json_loads(res))
            all_res.append(str(res) + "\n")
        except Exception as e:
            print(f"⚠️ Lỗi khi xử lý file {doc.get('file_name', 'Unknown')}: {e}")

    try:
        # with open("output.json", "w", encoding="UTF-8") as f:
        #     json.dump(all_res, f, ensure_ascii=False, indent=4)
        # print("✅ Đã lưu kết quả vào output.json")
        with open("output.txt", "w", encoding="utf-8") as f:
            f.writelines(all_res)
    except Exception as e:
        print(f"⚠️ Lỗi khi lưu file JSON: {e}")


# ========== TẠO FILE BÁO CÁO EXEL ==========
def create_excel_report(json_file="output.json", excel_file="output.xlsx"):
    sheet_name = "Kết quả chấm điểm"

    # Đọc JSON
    with open(json_file, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Chuyển dict 'detail' thành string để giữ trong 1 cột
    for d in data:
        d["detail"] = "\n".join([f"• {k}: {v}" for k, v in d["detail"].items()])

    #  Tạo DataFrame
    # Chứa tất cả các cột: name, total_point, detail, general
    df = pd.DataFrame(data)

    # Excel writer với xlsxwriter
    writer = pd.ExcelWriter(excel_file, engine="xlsxwriter")
    df.to_excel(writer, index=False, sheet_name=sheet_name)

    workbook = writer.book
    worksheet = writer.sheets[sheet_name]

    # Format wrap text + middle align
    cell_format_1 = workbook.add_format(
        {"align": "center", "valign": "vcenter", "text_wrap": True}
    )
    cell_format_2 = workbook.add_format({"valign": "vcenter", "text_wrap": True})

    # Áp dụng format + auto column width
    for i, col in enumerate(df.columns):
        max_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
        max_len = min(max_len, 80)  # giới hạn max width 80

        if i < 2:  # chỉ 2 cột đầu
            worksheet.set_column(i, i, max_len, cell_format_1)
        else:  # các cột còn lại
            worksheet.set_column(i, i, max_len, cell_format_2)

    # Lưu file
    writer.close()
    print("✅ Đã tạo file output.xlsx")


if __name__ == "__main__":
    grading()
    # create_excel_report()
    # print(read_topic_content())
