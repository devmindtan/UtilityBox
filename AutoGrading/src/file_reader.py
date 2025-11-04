from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
import fitz
from google.genai import types
from src.utils import detect_mine_type


class FileReader:
    def __init__(self, FOLDER_PATH, TOPIC_FILE_PATH):
        self.folder_path = Path(FOLDER_PATH)
        self.topic_file_path = Path(TOPIC_FILE_PATH)

    def read_multi_docx(self, max_files=None):
        data = []

        for i, doc_file in enumerate(self.folder_path.glob("*.docx"), start=1):
            if (max_files is not None and i > max_files):
                break

            doc = Document(doc_file)
            rels = doc.part.rels

            # ======= Đọc text =======
            all_text = []
            for para in doc.paragraphs:
                all_text.append(para.text.strip())
            text_output = "\n".join(all_text)

            # ======= Đọc ảnh =======
            image_parts = []
            for rel in rels.values():
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob
                    image_parts.append(
                        types.Part.from_bytes(
                            data=image_data, mime_type=detect_mine_type(image_data))
                    )

            data.append({
                "file_name": doc_file.stem.strip(),
                "text": text_output,
                "images": image_parts
            })

        return data

    def read_multi_pdf(self, max_files=None):
        data = []
        for pdf_file in self.folder_path.glob("*.pdf"):
            pdf = fitz.open(pdf_file)
            all_text = []
            image_parts = []

            for i, page in enumerate(pdf, start=1):
                if (max_files is not None and i > max_files):
                    break

                # Lấy text
                all_text.append(page.get_text())

                # Lấy ảnh
                for img in page.get_images(full=True):
                    xref = img[0]
                    image_data = pdf.extract_image(xref)["image"]
                    image_parts.append(types.Part.from_bytes(
                        data=image_data,
                        mime_type=detect_mine_type(image_data)
                    ))

            text_output = "\n".join(all_text)

            data.append({
                "file_name": pdf_file.stem.strip(),
                "text": text_output,
                "images": image_parts
            })

        return data

    def read_topic(self):
        file_path = Path(self.topic_file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File đề bài không tồn tại: {file_path}")

        # ===== HTML =====
        if file_path.suffix.lower() == ".html":
            with open(file_path, "r", encoding="utf-8") as f:
                html = f.read()
            soup = BeautifulSoup(html, "html.parser")
            return {"content": soup.get_text().strip(), "image_content": []}

        # ===== DOCX =====
        elif file_path.suffix.lower() == ".docx":
            doc = Document(file_path)
            content = "\n".join([p.text.strip()
                                for p in doc.paragraphs if p.text.strip()])
            rels = doc.part.rels.values()
            image_parts = []
            for rel in rels:
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob
                    image_parts.append(
                        types.Part.from_bytes(
                            data=image_data, mime_type=detect_mine_type(image_data))
                    )
            return {"content": content, "image_content": image_parts}

        else:
            raise ValueError("File đề bài phải là HTML hoặc DOCX")

    def read_all_content(self, max_file_docs=None, max_file_pdfs=None):
        all_data = []

        # DOCX
        # all_data.extend(self.read_multi_docx(max_file_docs))

        # PDF
        all_data.extend(self.read_multi_pdf(max_file_pdfs))

        return all_data


# file = r"C:\Documents\Code\UtilityBox\AutoGrading\assignments"
# file_path = r"C:\Documents\Code\UtilityBox\AutoGrading\assignments\topic\baitap_lab5.html"
# fr = FileReader(file, file_path)
# print(fr.read_all_content())
