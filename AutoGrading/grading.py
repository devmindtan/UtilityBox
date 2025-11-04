import re
from docx import Document
from bs4 import BeautifulSoup
from google import genai
from google.genai import types
from pathlib import Path
import json
import pandas as pd
import filetype
import fitz
import os
from dotenv import load_dotenv

load_dotenv()

FOLDER_PATH = r"C:\Documents\Code\UtilityBox\AutoGrading\Tuan 6"
TOPIC_FILE = "Lab2.CacWidgetCoBan.docx"
API_KEY = os.getenv('API_KEY')
MODEL_NAME = "gemini-2.5-flash"
# RUBRIC = """
# Rubric chấm điểm (tổng 10 điểm):
# 1. Nộp code (4đ)
# 2. Giải thích ý tưởng (4đ) (Giải thích ý tưởng của mình trước khi code không giải thích những dòng code)
#     - Ngắn gọn xúc tích
#     - Thể hiện mình hiểu được bài tập
#     - Đánh giá cao những ý tưởng cá nhân không theo lối mòn
#     => Nếu không làm được những cái trên sẽ bị trừ điểm
# 3. Có kết chạy được (1đ)
# 4. Trình bày gọn gàng đầy đủ (1đ)
# Đặc biệt: nếu tên file không đúng theo quy chuẩn này HoTen_MSSV_lab6 (Ví dụ: NguyenVanA_22133422_lab6) sẽ trừ 1 điểm
# """

RUBRIC = """
Rubric chấm điểm (tổng 10 điểm):
Bài tập lần này không yêu cầu gì nhiều chỉ cần code theo xong thì giải thích những gì bạn hiểu trong đoạn code mẫu.
1. Code (7đ)
2. Giải thích (3đ)
Lưu ý: giải thích bằng chính văn của mình, không đạo, không copy, hiểu gì nói đó, nếu vi phạm thì sẽ bị trừ điểm
"""


# ========== INIT CLIENT ==========
client = genai.Client(api_key=API_KEY)


# Xác định kiểu của ảnh
def detect_mine_type(image_data):
    mime_type = ""

    kind = filetype.guess(image_data)
    if kind:
        mime_type = kind.mime
    else:
        mime_type = "image/png"

    return mime_type


# Xử lí đầu vào
def safe_json_loads(text):
    # Cắt bỏ mọi thứ ngoài JSON thật
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        clean = match.group(0).strip()
        return json.loads(clean)
    else:
        raise ValueError("Không tìm thấy JSON hợp lệ trong phản hồi")


def read_doc_content(max_files=None):
    folder_path = Path(FOLDER_PATH)
    data = []

    for i, doc_file in enumerate(folder_path.glob("*.docx"), start=1):
        if (max_files is not None and i > max_files):
            break

        doc = Document(doc_file)
        rels = doc.part.rels

        # ======= Đọc text =======
        all_text = []
        for para in doc.paragraphs:
            all_text.append(para.text.strip())
        text_output = "\n".join(all_text)

        # ======= Đọc ảnh =======
        image_parts = []
        for rel in rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                image_parts.append(
                    types.Part.from_bytes(
                        data=image_data, mime_type=detect_mine_type(image_data))
                )

        data.append({
            "file_name": doc_file.stem.strip(),
            "text": text_output,
            "images": image_parts
        })

    return data


def read_pdf_content():
    folder_path = Path(FOLDER_PATH)
    data = []

    for pdf_file in folder_path.glob("*.pdf"):
        pdf = fitz.open(pdf_file)
        all_text = []
        image_parts = []

        for page in pdf:
            # Lấy text
            all_text.append(page.get_text())

            # Lấy ảnh
            for img in page.get_images(full=True):
                xref = img[0]
                image_data = pdf.extract_image(xref)["image"]
                image_parts.append(types.Part.from_bytes(
                    data=image_data,
                    mime_type=detect_mine_type(image_data)
                ))

        text_output = "\n".join(all_text)

        data.append({
            "file_name": pdf_file.stem.strip(),
            "text": text_output,
            "images": image_parts
        })

    return data


def read_topic_content():
    path = Path(TOPIC_FILE)
    if not path.exists():
        raise FileNotFoundError(f"File đề bài không tồn tại: {TOPIC_FILE}")

    # ===== HTML =====
    if path.suffix.lower() == ".html":
        with open(path, "r", encoding="utf-8") as f:
            html = f.read()
        soup = BeautifulSoup(html, "html.parser")
        return {"content": soup.get_text().strip(), "image_content": []}

    # ===== DOCX =====
    elif path.suffix.lower() == ".docx":
        doc = Document(path)
        content = "\n".join([p.text.strip()
                             for p in doc.paragraphs if p.text.strip()])
        rels = doc.part.rels.values()
        image_parts = []
        for rel in rels:
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                image_parts.append(
                    types.Part.from_bytes(
                        data=image_data, mime_type=detect_mine_type(image_data))
                )
        return {"content": content, "image_content": image_parts}

    else:
        raise ValueError("File đề bài phải là HTML hoặc DOCX")


def read_all_content():
    all_data = []

    # DOCX
    all_data.extend(read_doc_content())

    # PDF
    all_data.extend(read_pdf_content())

    return all_data


# ========== TẠO PROMPT ==========
def create_prompt(file_name, content, level, response):
    return f"""
    Bạn là giảng viên đại học chấm bài sinh viên. Hãy đọc kỹ **đề bài**, **rubric**, và **nội dung bài làm** (text hoặc hình ảnh).

    ## Nhiệm vụ:
    1. Đánh giá chi tiết từng tiêu chí trong rubric.
    2. Ghi rõ điểm cho từng tiêu chí.
    3. Tính tổng điểm (thang 10) (làm tròn .5 trở lên là lên, dưới .5 là xuống).
    4. Nhận xét ngắn gọn (dưới 30 chữ/tiêu chí), chuyên nghiệp, xưng "em".
    5. Khi viết nhận xét chi tiết và tổng quát, **tham khảo `response` trong `ai_review` để điều chỉnh cách diễn đạt**, ví dụ:
       - Nếu `response` nhấn mạnh bài vẫn trong giới hạn cho phép, hãy dùng lời khích lệ, nhẹ nhàng.
       - Nếu `response` nhấn mạnh bài có dấu hiệu AI rõ rệt, hãy thể hiện sự nhắc nhở/nhấn mạnh nhưng vẫn trung thực, chuyên nghiệp.
    6. Trả về **duy nhất** một JSON hợp lệ theo mẫu sau:

    ---
    {{
      "name": "{file_name}",
      "total_point": <số điểm trên 10>,
      "detail": {{
        "Tên tiêu chí 1": "[Điểm] điểm — [Nhận xét ngắn dựa trên response]",
        "Tên tiêu chí 2": "[Điểm] điểm — [Nhận xét ngắn dựa trên response]",
        ...
      }},
      "general": "<nhận xét tổng quát 1–2 câu dựa trên response>",
      "ai_review": {{
         "muc_do": {level},
         "phan_hoi": "{response}"
      }}
    }}
    ---

    ## Thông tin chấm:
    **Đề bài:**
    {read_topic_content()}

    **Rubric:**
    {RUBRIC}

    **Nội dung bài làm (text):**
    {content}

    Nếu có hình ảnh, hãy xét nội dung trong ảnh nữa.
    """


# ========== GỌI GEMINI ==========
def genemi_call(prompt_text, images):
    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=[
            types.Part.from_text(text=prompt_text),
            *images
        ]
    )

    return response.text


# ========== CHẤM ĐIỂM HÀNG LOẠT ==========
def grading():
    all_res = []
    for doc in read_all_content():
        print(f"📄 {doc["file_name"]} — Đã đọc {len(doc["images"])} ảnh")
        file_name = doc["file_name"]
        content = doc["text"]
        images = doc["images"]
        prompt_text = create_prompt(file_name, content)
        res = genemi_call(prompt_text, images)
        all_res.append(safe_json_loads(res))

    with open("output.json", "w", encoding="UTF-8") as f:
        json.dump(all_res, f, ensure_ascii=False, indent=4)


# ========== TẠO FILE BÁO CÁO EXEL ==========
def create_excel_report(json_file="output.json", excel_file="output.xlsx"):
    sheet_name = "Kết quả chấm điểm"

    # Đọc JSON
    with open(json_file, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Chuyển dict 'detail' thành string để giữ trong 1 cột
    for d in data:
        d['detail'] = "\n".join(
            [f"• {k}: {v}" for k, v in d['detail'].items()])

    #  Tạo DataFrame
    # Chứa tất cả các cột: name, total_point, detail, general
    df = pd.DataFrame(data)

    # Excel writer với xlsxwriter
    writer = pd.ExcelWriter(excel_file, engine='xlsxwriter')
    df.to_excel(writer, index=False, sheet_name=sheet_name)

    workbook = writer.book
    worksheet = writer.sheets[sheet_name]

    # Format wrap text + middle align
    cell_format_1 = workbook.add_format({
        'align': 'center',
        'valign': 'vcenter',
        'text_wrap': True
    })
    cell_format_2 = workbook.add_format({
        'valign': 'vcenter',
        'text_wrap': True
    })

    # Áp dụng format + auto column width
    for i, col in enumerate(df.columns):
        max_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
        max_len = min(max_len, 80)  # giới hạn max width 80

        if i < 2:  # chỉ 2 cột đầu
            worksheet.set_column(i, i, max_len, cell_format_1)
        else:  # các cột còn lại
            worksheet.set_column(i, i, max_len, cell_format_2)

    # Lưu file
    writer.close()
    print("✅ Đã tạo file output.xlsx")


if __name__ == "__main__":
    grading()
    create_excel_report()
    # print(read_topic_content())
